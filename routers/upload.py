from fastapi import APIRouter, UploadFile, File, Depends, HTTPException
from sqlalchemy.orm import Session
import openpyxl
import docx
from datetime import datetime, date
from sqlalchemy import String
import models
from database import get_db
from auth import get_current_user

router = APIRouter(prefix="/upload", tags=["upload"])

HEADER_MAP = {
    "file code": "file_code",
    "employee code": "file_code",
    "staff no": "file_code",
    "project": "project",
    "full name": "full_name",
    "name": "full_name",
    "status": "status",
    "staff status": "status",
    "employment status": "status",
    "position": "position",
    "designation": "position",
    "location": "location",
    "site": "location",
    "contact": "contact_number",
    "contact number": "contact_number",
    "phone": "contact_number",
    "locker": "locker",
    "date of appointment": "date_of_appointment",
    "appointment date": "date_of_appointment",
    "contract start": "contract_start",
    "contract end": "contract_end",
    "contract review date": "contract_review_date",
    "review date": "contract_review_date",
    "probation end": "probation_end",
    "notice period": "notice_period",
    "employment type": "employment_type",
    "contract type": "employment_type",
    "missing application/resume": "missing_app_resume",
    "missing app resume": "missing_app_resume",
    "missing appointment letter": "missing_appointment_letter",
    "missing academic docs": "missing_academic_docs",
    "missing recruitment notes": "missing_recruitment_notes",
    "missing staff id form": "missing_staff_id_form",
    "missing performance appraisals": "missing_performance_appraisals",
    "missing national id": "missing_national_id",
    "missing policy declaration": "missing_policy_declaration",
    "missing end of contract notice": "missing_end_of_contract_notice",
    # photo / avatar mapping (allows Excel to include image filename or URL)
    "photo": "photo_url",
    "photo url": "photo_url",
    "image": "photo_url",
    "picture": "photo_url",
    "profile image": "photo_url",
    "profile photo": "photo_url",
    "avatar": "photo_url",
}

BOOL_KEYS = {
    "yes",
    "true",
    "1",
    "x",
    "y",
    "present",
    "missing",
}

@router.post("/excel")
async def upload_excel(file: UploadFile = File(...), db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    try:
        if current_user.role != "hr_admin":
            raise HTTPException(status_code=403, detail="Only HR admin can upload")

        if not file.filename.endswith(('.xlsx', '.xls')):
            raise HTTPException(status_code=400, detail="File must be Excel format (.xlsx or .xls)")

        workbook = openpyxl.load_workbook(file.file)
        imported = 0
        updated = 0
        errors = []
        sheets_processed = []
        
        for sheet_index, sheet in enumerate(workbook.worksheets, 1):
            print(f"\nProcessing sheet {sheet_index}: '{sheet.title}'")
            header_row_index, headers = _find_headers(sheet)
            if not headers:
                print(f"  No headers found in sheet '{sheet.title}', skipping...")
                continue
            
            sheets_processed.append(sheet.title)
            for row in sheet.iter_rows(min_row=header_row_index + 1, values_only=True):
                try:
                    row_data = _map_row_data(headers, row)
                    if not row_data.get("file_code") or not row_data.get("full_name"):
                        continue
                    row_data = _normalize_row_data(row_data)
                    print(f"Processing: {row_data.get('full_name')} (file_code: {row_data.get('file_code')}, status: {row_data.get('status')}, project: {row_data.get('project')})")
                    
                    # Ensure contact_number is always a string
                    if 'contact_number' in row_data and row_data['contact_number'] is not None:
                        row_data['contact_number'] = str(row_data['contact_number'])
                    
                    existing = db.query(models.Employee).filter(models.Employee.file_code == row_data["file_code"]).first()
                    if existing:
                        # Update existing record
                        for key, value in row_data.items():
                            if hasattr(existing, key) and value is not None:
                                setattr(existing, key, value)
                        updated += 1
                        print(f"Updated existing employee: {row_data.get('full_name')}")
                    else:
                        # Create new record
                        db.add(models.Employee(**row_data))
                        imported += 1
                        print(f"Added new employee: {row_data.get('full_name')}")
                    
                    # Commit after each successful operation to avoid batch conflicts
                    db.commit()
                    
                except Exception as row_error:
                    print(f"Error processing employee {row_data.get('full_name', 'Unknown')}: {row_error}")
                    errors.append(f"Row {row_data.get('full_name', 'Unknown')}: {str(row_error)}")
                    db.rollback()  # Rollback this specific row's changes
                    continue

        if errors:
            message = f"Processed {len(sheets_processed)} sheets: {', '.join(sheets_processed)}. Imported {imported} new records, updated {updated} existing records. {len(errors)} errors occurred."
            return {"message": message, "errors": errors[:10]}
        else:
            message = f"Processed {len(sheets_processed)} sheets: {', '.join(sheets_processed)}. Successfully imported {imported} new records and updated {updated} existing records"
            return {"message": message}
    except HTTPException as e:
        raise e
    except Exception as e:
        print(f"Upload error: {e}")
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")


def _find_headers(sheet):
    for row_index, row in enumerate(sheet.iter_rows(min_row=1, max_row=12, values_only=True), start=1):
        if not row:
            continue
        normalized = [str(cell).strip().lower() if cell is not None else "" for cell in row]
        if any(key in normalized for key in ("file code", "full name", "project", "contact")):
            return row_index, [HEADER_MAP.get(cell.strip().lower(), None) if cell else None for cell in normalized]
    return None, []


def _map_row_data(headers, row):
    data = {}
    for idx, key in enumerate(headers):
        if key:
            data[key] = row[idx] if idx < len(row) else None
    return data


def _normalize_row_data(data):
    normalized = {}
    for key, value in data.items():
        if value is None:
            normalized[key] = None
            continue
        if key in {"missing_app_resume", "missing_appointment_letter", "missing_academic_docs", "missing_recruitment_notes", "missing_staff_id_form", "missing_performance_appraisals", "missing_national_id", "missing_policy_declaration", "missing_end_of_contract_notice"}:
            normalized[key] = _parse_bool(value)
        elif key in {"date_of_appointment", "contract_start", "contract_end", "contract_review_date", "probation_end"}:
            normalized[key] = _parse_date(value)
        else:
            normalized[key] = str(value).strip() if isinstance(value, str) else value
    
    # Infer status from project name if status is empty
    if not normalized.get("status") and normalized.get("project"):
        project = str(normalized.get("project", "")).upper()
        if "EXITED" in project or "EXIT" in project:
            normalized["status"] = "Exited"
        elif "RECESS" in project or "ON RECESS" in project:
            normalized["status"] = "On Recess"
        else:
            # Default to Active if status not specified and project doesn't indicate otherwise
            normalized["status"] = "Active"
    elif not normalized.get("status"):
        # If no project either, default to Active
        normalized["status"] = "Active"
    
    return normalized


def _parse_bool(value):
    if isinstance(value, bool):
        return value
    text = str(value).strip().lower()
    return text in BOOL_KEYS


def _parse_date(val):
    if isinstance(val, datetime):
        return val.date()
    if isinstance(val, str):
        text = val.strip()
        for fmt in ["%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%m/%d/%Y", "%m-%d-%Y"]:
            try:
                return datetime.strptime(text, fmt).date()
            except ValueError:
                continue
    return None


def _extract_docx_fields(document):
    fields = {}

    for table in document.tables:
        for row in table.rows:
            if len(row.cells) >= 2:
                key = row.cells[0].text.strip().lower()
                value = row.cells[1].text.strip()
                if key:
                    fields[key] = value

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if ':' in text:
            parts = text.split(':', 1)
            key = parts[0].strip().lower()
            value = parts[1].strip()
            if key:
                fields[key] = value

    return fields


def _map_docx_data(field_map):
    data = {}
    for raw_key, value in field_map.items():
        normalized_key = raw_key.strip().lower()
        if normalized_key in HEADER_MAP:
            data[HEADER_MAP[normalized_key]] = value
            continue

        for known_key, mapped_key in HEADER_MAP.items():
            if known_key in normalized_key:
                data[mapped_key] = value
                break

        if 'employee id' in normalized_key and 'file_code' not in data:
            data['file_code'] = value
        if 'employee_name' in normalized_key and 'full_name' not in data:
            data['full_name'] = value
        if 'full name' in normalized_key and 'full_name' not in data:
            data['full_name'] = value

    return data


@router.post("/word")
async def upload_word(file: UploadFile = File(...), db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    try:
        if current_user.role != "hr_admin":
            raise HTTPException(status_code=403, detail="Only HR admin can upload Word documents")

        if not file.filename.lower().endswith('.docx'):
            raise HTTPException(status_code=400, detail="File must be Word document format (.docx)")

        document = docx.Document(file.file)
        raw_fields = _extract_docx_fields(document)
        if not raw_fields:
            raise HTTPException(status_code=400, detail="No fields could be extracted from the Word document")

        row_data = _map_docx_data(raw_fields)
        if not row_data.get('file_code') or not row_data.get('full_name'):
            raise HTTPException(status_code=400, detail='Word document must include employee file code and full name')

        row_data = _normalize_row_data(row_data)
        if 'contact_number' in row_data and row_data['contact_number'] is not None:
            row_data['contact_number'] = str(row_data['contact_number'])

        existing = db.query(models.Employee).filter(models.Employee.file_code == row_data['file_code']).first()
        if existing:
            for key, value in row_data.items():
                if hasattr(existing, key) and value is not None:
                    setattr(existing, key, value)
            db.commit()
            return {"message": f"Updated employee {row_data.get('full_name')} from Word document"}

        db.add(models.Employee(**row_data))
        db.commit()
        return {"message": f"Imported employee {row_data.get('full_name')} from Word document"}
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Word upload failed: {str(e)}")


@router.post("/payslips_excel")
async def upload_payslips_excel(file: UploadFile = File(...), db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    """Import payslips from an Excel file. Expected headers: file_code or employee_id, period_start, period_end, gross_pay, tax, deductions, pdf_url (optional)"""
    try:
        if current_user.role not in ("hr_admin", "project_manager"):
            raise HTTPException(status_code=403, detail="Insufficient permissions to upload payslips")

        if not file.filename.endswith(('.xlsx', '.xls')):
            raise HTTPException(status_code=400, detail="File must be Excel format (.xlsx or .xls)")

        workbook = openpyxl.load_workbook(file.file)
        imported = 0
        errors = []
        for sheet in workbook.worksheets:
            header_row_index = None
            headers = None
            # find header row
            for row_index, row in enumerate(sheet.iter_rows(min_row=1, max_row=12, values_only=True), start=1):
                if not row:
                    continue
                normalized = [str(cell).strip().lower() if cell is not None else "" for cell in row]
                if any(k in normalized for k in ("file code", "employee code", "employee id", "period start", "gross", "gross_pay", "net_pay", "period_end", "period end")):
                    header_row_index = row_index
                    headers = [cell.strip().lower() if cell else "" for cell in row]
                    break
            if not headers:
                continue

            # normalize header mapping
            header_map = []
            for h in headers:
                h = h.lower()
                if "file" in h and ("code" in h or "employee code" in h):
                    header_map.append('file_code')
                elif "employee id" in h or h == 'employee_id':
                    header_map.append('employee_id')
                elif "period start" in h or "start" in h:
                    header_map.append('period_start')
                elif "period end" in h or "end" in h:
                    header_map.append('period_end')
                elif "gross" in h:
                    header_map.append('gross_pay')
                elif "tax" in h:
                    header_map.append('tax')
                elif "deduct" in h:
                    header_map.append('deductions')
                elif "net" in h:
                    header_map.append('net_pay')
                elif "pdf" in h or "url" in h or "file url" in h:
                    header_map.append('pdf_url')
                else:
                    header_map.append(h.replace(' ', '_'))

            for row in sheet.iter_rows(min_row=header_row_index + 1, values_only=True):
                try:
                    row_dict = {}
                    for idx, key in enumerate(header_map):
                        if idx < len(row):
                            row_dict[key] = row[idx]

                    # resolve employee id
                    emp_id = None
                    if row_dict.get('employee_id'):
                        try:
                            emp_id = int(row_dict.get('employee_id'))
                        except Exception:
                            emp_id = None
                    if not emp_id and row_dict.get('file_code'):
                        emp = db.query(models.Employee).filter(models.Employee.file_code == str(row_dict.get('file_code'))).first()
                        if emp:
                            emp_id = emp.id

                    if not emp_id:
                        errors.append(f"Row missing employee identifier: {row_dict}")
                        continue

                    period_start = _parse_date(row_dict.get('period_start'))
                    period_end = _parse_date(row_dict.get('period_end'))
                    try:
                        gross = float(row_dict.get('gross_pay') or 0)
                    except Exception:
                        gross = 0.0
                    try:
                        tax = float(row_dict.get('tax') or 0)
                    except Exception:
                        tax = 0.0
                    try:
                        deductions = float(row_dict.get('deductions') or 0)
                    except Exception:
                        deductions = 0.0

                    net = float(row_dict.get('net_pay') or (gross - tax - deductions))

                    payslip = models.Payslip(
                        employee_id=emp_id,
                        period_start=period_start,
                        period_end=period_end,
                        gross_pay=gross,
                        tax=tax,
                        deductions=deductions,
                        net_pay=net,
                        pdf_url=row_dict.get('pdf_url'),
                        generated_by=current_user.id
                    )
                    db.add(payslip)
                    db.commit()
                    imported += 1
                except Exception as e:
                    db.rollback()
                    errors.append(str(e))
                    continue

        return {"imported": imported, "errors": errors[:20]}
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))
